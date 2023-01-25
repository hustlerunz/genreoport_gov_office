from tkinter import N
from tkinter.messagebox import NO
import docx
import pandas as pd
from docx import Document
import openpyxl
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.enum.text import WD_UNDERLINE
import configparser
import sys 
from docx.shared import RGBColor

config = configparser.ConfigParser()
config.read('go.ini')
source_file_data = config['file']['source_data']
master_file = config['file']['form_master']
row_run = int(config['file']['ro_start'])
ch = 0

book = openpyxl.load_workbook(source_file_data)
sheet = book.active
row_count = sheet.max_row

sara = ['่','้','๊','๋','ิ','ี','ื','ุ','ู','์','ั']
sara2 = ['โ','ไ','เ','ใ','.']
#print(sara)

print("excel_sum_row = ",row_count)  
print("========run========")  
for x in range(row_count-1):
    chek_len_count = 0 #check char
    chek_len_count2 = 0 #check char
    var2_chk_half = 0
    var1_chk_half = 0
    #print(x)
    chk_name = sheet.cell(row=row_run+x, column=3).value
    if chk_name == None:
        book.close()
        sys.exit()

    doc = Document(master_file)
    name = sheet.cell(row=row_run+x, column=3).value
    sname = sheet.cell(row=row_run+x, column=4).value
    position = sheet.cell(row=row_run+x, column=5).value
    no_position = sheet.cell(row=row_run+x, column=6).value
    own = sheet.cell(row=row_run+x, column=7).value

    st_point = sheet.cell(row=row_run+x,column=8).value
    st_point2 = sheet.cell(row=row_run+x,column=9).value

    l1_sick = sheet.cell(row=row_run+x,column=10).value
    if l1_sick==None:
        l1_sick = ""
    l1_bus = sheet.cell(row=row_run+x,column=11).value
    if l1_bus==None:
        l1_bus = ""
    l1_ordan = sheet.cell(row=row_run+x,column=12).value
    if l1_ordan==None:
        l1_ordan = ""
    l1_giv = sheet.cell(row=row_run+x,column=13).value
    if l1_giv==None:
        l1_giv = ""	
    l1_late = sheet.cell(row=row_run+x,column=14).value
    if l1_late==None:
        l1_late = ""
    l1_los = sheet.cell(row=row_run+x,column=15).value
    if l1_los==None:
        l1_los = ""
    l1_sum = sheet.cell(row=row_run+x,column=16).value
    if l1_sum==None:
        l1_sum = ""
    l2_sick = sheet.cell(row=row_run+x,column=17).value
    if l2_sick==None:
        l2_sick = ""
    l2_bus = sheet.cell(row=row_run+x,column=18).value
    if l2_bus==None:
        l2_bus = ""
    l2_ordan = sheet.cell(row=row_run+x,column=19).value
    if l2_ordan==None:
        l2_ordan = ""
    l2_giv = sheet.cell(row=row_run+x,column=20).value
    if l2_giv==None:
        l2_giv = ""
    l2_late = sheet.cell(row=row_run+x,column=21).value
    if l2_late==None:
        l2_late = ""
    l2_los = sheet.cell(row=row_run+x,column=22).value
    if l2_los==None:
        l2_los = ""
    l2_sum = sheet.cell(row=row_run+x,column=23).value
    if l2_sum==None:
        l2_sum = ""
    discipline1_1 = sheet.cell(row=row_run+x,column=24).value
    if discipline1_1==None:
        discipline1_1 = ""
    discipline1_2 = sheet.cell(row=row_run+x,column=25).value
    if discipline1_2==None:
        discipline1_2 = ""
    discipline1_3 = sheet.cell(row=row_run+x,column=26).value
    if discipline1_3==None:
        discipline1_3 = ""
    discipline1_4 = sheet.cell(row=row_run+x,column=27).value
    if discipline1_4==None:
        discipline1_4 = ""
    discipline1_5 = sheet.cell(row=row_run+x,column=28).value
    if discipline1_5==None:
        discipline1_5 = ""
    discipline1_6 = sheet.cell(row=row_run+x,column=29).value
    if discipline1_6==None:
        discipline1_6 = ""
    discipline2_1 = sheet.cell(row=row_run+x,column=30).value
    if discipline2_1==None:
        discipline2_1 = ""
    discipline2_2 = sheet.cell(row=row_run+x,column=31).value
    if discipline2_2==None:
        discipline2_2 = ""
    discipline2_3 = sheet.cell(row=row_run+x,column=32).value
    if discipline2_3==None:
        discipline2_3 = ""
    discipline2_4 = sheet.cell(row=row_run+x,column=33).value
    if discipline2_4==None:
        discipline2_4 = ""
    discipline2_5 = sheet.cell(row=row_run+x,column=34).value
    if discipline2_5==None:
        discipline2_5 = ""
    discipline2_6 = sheet.cell(row=row_run+x,column=35).value
    if discipline2_6==None:
        discipline2_6 = ""

    var1 = '   '+name+'  '+sname+'   ' 
    var2 = '   '+position+'   '
    var3 = ' '
    var33 = '.'
    var4 = '  '+no_position+'  '
    var5 = '  '+own+'  '

    chk_len = var1+var2
    chk_len2 = var4+var5
    for cha in chk_len:
        for chk2 in sara:
            if cha == chk2:
                chek_len_count = chek_len_count+1
        for chk_h in sara2:
            if cha == chk_h:
                var1_chk_half = var1_chk_half+1        
    for cha2 in chk_len2:
        for chk2 in sara:
            if cha2 == chk2:
                chek_len_count2 = chek_len_count2+1
        for chk_h in sara2:
            if cha2 == chk_h:
                var2_chk_half = var2_chk_half+1     
    print("var1_lenth = ",len(chk_len))
    print("var2_lenth = ",len(chk_len2))                        
    print("var1_upper_under = ",chek_len_count)            
    print("var2_upper_under = ",chek_len_count2)
    print("var1_half_c = ",var1_chk_half)            
    print("var2_half_c = ",var2_chk_half)                 
    
    print("var1_sum_postion = ",((len(chk_len)-chek_len_count)-chek_len_count2/2))
    print("var1_sum_paragraph = ",19.5+(len(chk_len)-chek_len_count)-(chek_len_count2/2))

    print("var2_sum_postion = ",(len(chk_len2)-chek_len_count2)-(chek_len_count2/2))
    print("var2_sum_paragraph = ",22+(len(chk_len2)-chek_len_count2)-(chek_len_count2/2))

    sp_num = 72-(len(chk_len)-chek_len_count)-(chek_len_count2/2)
    #sp_num = 55-(len(chk_len)-chek_len_count)
    sp_num2 = 66-(len(chk_len2)-chek_len_count2)-(chek_len_count2/2)
    print("sp_num = ",sp_num)
    print("sp_num2 = ",sp_num2)
    def space_len(var_len,pp_p):
        for x in range(var_len):
            v11 = doc.paragraphs[pp_p].add_run(var3)
            v11.font.name = 'TH SarabunIT๙'
            v11.font.size = Pt(16)
            v11.underline = True
            v11.underline = WD_UNDERLINE.DOTTED

    newdata = []

    var = doc.paragraphs[4].text = 'ชื่อผู้รับการประเมิน'
    v2 = doc.paragraphs[4].add_run(var1)
    v2.font.name = 'TH SarabunIT๙'
    v2.font.size = Pt(16)
    v2.underline = True
    v2.underline = WD_UNDERLINE.DOTTED
    v3 = doc.paragraphs[4].add_run('ตำแหน่ง')
    v3.font.name = 'TH SarabunIT๙'
    v3.font.size = Pt(16)
    v4 = doc.paragraphs[4].add_run(var2)
    v4.font.name = 'TH SarabunIT๙'
    v4.font.size = Pt(16)
    v4.underline = True
    v4.underline = WD_UNDERLINE.DOTTED
    
    space_len(int(sp_num),4)
    v66 = doc.paragraphs[4].add_run(var33)
    v66.font.name = 'TH SarabunIT๙'
    v66.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    v66.font.size = Pt(16)
    v66.underline = True
    v66.underline = WD_UNDERLINE.DOTTED

    #v5 = doc.paragraphs[5].add_run('กลุ่มงาน')
    v5 = doc.paragraphs[5].text = 'กลุ่มงาน'
    #v5.font.name = 'TH SarabunIT๙'
    #v5.font.size = Pt(16)

    v6 = doc.paragraphs[5].add_run(var3)
    v6.font.name = 'TH SarabunIT๙'
    v6.font.size = Pt(16)
    v6.underline = True
    v6.underline = WD_UNDERLINE.DOTTED

    v6_6 = doc.paragraphs[5].add_run(' - ')
    v6_6.font.name = 'TH SarabunIT๙'
    v6_6.font.size = Pt(16)
    v6_6.underline = True
    v6_6.underline = WD_UNDERLINE.DOTTED

    v666 = doc.paragraphs[5].add_run(var3)
    v666.font.name = 'TH SarabunIT๙'
    v666.font.size = Pt(16)
    v666.underline = True
    v666.underline = WD_UNDERLINE.DOTTED

    #v7 = doc.paragraphs[5].text = 'เลขที่ตำแหน่ง'
    v7 = doc.paragraphs[5].add_run('เลขที่ตำแหน่ง')
    v7.font.name = 'TH SarabunIT๙'
    v7.font.size = Pt(16)
    v8 = doc.paragraphs[5].add_run(var4)
    v8.font.name = 'TH SarabunIT๙'
    v8.font.size = Pt(16)
    v8.underline = True
    v8.underline = WD_UNDERLINE.DOTTED
    v9 = doc.paragraphs[5].add_run('สังกัด')
    v9.font.name = 'TH SarabunIT๙'
    v9.font.size = Pt(16)
    v10 = doc.paragraphs[5].add_run(var5)
    v10.font.name = 'TH SarabunIT๙'
    v10.font.size = Pt(16)
    v10.underline = True
    v10.underline = WD_UNDERLINE.DOTTED

    space_len(int(sp_num2),5)
    
    v111 = doc.paragraphs[5].add_run(var33)
    v111.font.name = 'TH SarabunIT๙'
    v111.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    v111.font.size = Pt(16)
    v111.underline = True
    v111.underline = WD_UNDERLINE.DOTTED
    #data Round 1
    if int(st_point) >= 90: 
        v12 = doc.tables[0].cell(1,1).paragraphs[0].text = str(st_point)
        doc.tables[0].cell(1,1).paragraphs[0].alignment = 1
        newdata.append(doc.tables[0].cell(1,1).paragraphs[0])
    elif st_point >= 80 and st_point <=89:
        v12 = doc.tables[0].cell(1,2).paragraphs[0].text = str(st_point)
        doc.tables[0].cell(1,2).paragraphs[0].alignment = 1
        newdata.append(doc.tables[0].cell(1,2).paragraphs[0])
    elif int(st_point) >= 70 and int(st_point) <=79:
        v12 = doc.tables[0].cell(1,3).paragraphs[0].text = str(st_point)
        doc.tables[0].cell(1,3).paragraphs[0].alignment = 1
        newdata.append(doc.tables[0].cell(1,3).paragraphs[0])
    elif int(st_point) >= 60 and int(st_point) <=69:
        v12 = doc.tables[0].cell(1,4).paragraphs[0].text = str(st_point)
        doc.tables[0].cell(1,4).paragraphs[0].alignment = 1
        newdata.append(doc.tables[0].cell(1,4).paragraphs[0])
    else:
        v12 = doc.tables[0].cell(1,5).paragraphs[0].text = str(st_point)
        doc.tables[0].cell(1,5).paragraphs[0].alignment = 1
        newdata.append(doc.tables[0].cell(1,5).paragraphs[0])
    #round2          
    if st_point2 != None:
        if int(st_point2) >= 90: 
            v2_12 = doc.tables[0].cell(2,1).paragraphs[0].text = str(st_point2)
            doc.tables[0].cell(2,1).paragraphs[0].alignment = 1
            newdata.append(doc.tables[0].cell(2,1).paragraphs[0])
        
        elif st_point2 >= 80 and st_point2 <=89:
            v2_12 = doc.tables[0].cell(2,2).paragraphs[0].text = str(st_point2)
            doc.tables[0].cell(2,2).paragraphs[0].alignment = 1
            newdata.append(doc.tables[0].cell(2,2).paragraphs[0])
        elif int(st_point2) >= 70 and int(st_point2) <=79:
            v2_12 = doc.tables[0].cell(2,3).paragraphs[0].text = str(st_point2)
            doc.tables[0].cell(2,3).paragraphs[0].alignment = 1
            newdata.append(doc.tables[0].cell(2,3).paragraphs[0])
        elif int(st_point2) >= 60 and int(st_point2) <=69:
            v2_12 = doc.tables[0].cell(2,4).paragraphs[0].text = str(st_point2)
            doc.tables[0].cell(2,4).paragraphs[0].alignment = 1
            newdata.append(doc.tables[0].cell(2,4).paragraphs[0])
        else:
            v2_12 = doc.tables[0].cell(2,5).paragraphs[0].text = str(st_point2)
            doc.tables[0].cell(2,5).paragraphs[0].alignment = 1
            newdata.append(doc.tables[0].cell(2,5).paragraphs[0]) 

    v13 = doc.tables[1].cell(1,1).paragraphs[0].text = str(l1_sick)
    v14 = doc.tables[1].cell(1,2).paragraphs[0].text = str(l1_bus)
    v15 = doc.tables[1].cell(1,3).paragraphs[0].text = str(l1_ordan)
    v16 = doc.tables[1].cell(1,4).paragraphs[0].text = str(l1_giv)
    v17 = doc.tables[1].cell(1,5).paragraphs[0].text = str(l1_late)
    v18 = doc.tables[1].cell(1,6).paragraphs[0].text = str(l1_los)
    v19 = doc.tables[1].cell(1,7).paragraphs[0].text = str(l1_sum)

    v20 = doc.tables[1].cell(2,1).paragraphs[0].text = str(l2_sick)
    v21 = doc.tables[1].cell(2,2).paragraphs[0].text = str(l2_bus)
    v22 = doc.tables[1].cell(2,3).paragraphs[0].text = str(l2_ordan)
    v23 = doc.tables[1].cell(2,4).paragraphs[0].text = str(l2_giv)
    v24 = doc.tables[1].cell(2,5).paragraphs[0].text = str(l2_late)
    v25 = doc.tables[1].cell(2,6).paragraphs[0].text = str(l2_los)
    v26 = doc.tables[1].cell(2,7).paragraphs[0].text = str(l2_sum)

    #v27	= doc.tables[2].cell(1,1).paragraphs[0].text =	str(l2_sick)
    #v28	= doc.tables[2].cell(1,2).paragraphs[0].text =	str(l2_bus)
    #v29	= doc.tables[2].cell(1,3).paragraphs[0].text =	str(l2_ordan)
    #v30	= doc.tables[2].cell(1,4).paragraphs[0].text =	str(l2_giv)
    #v31	= doc.tables[2].cell(1,5).paragraphs[0].text =	str(l2_late)
    #v32	= doc.tables[2].cell(1,6).paragraphs[0].text =	str(l2_los)

    v33	= doc.tables[2].cell(1,1).paragraphs[0].text =	str(discipline1_1)
    v32	= doc.tables[2].cell(1,2).paragraphs[0].text =	str(discipline1_2)
    v33	= doc.tables[2].cell(1,3).paragraphs[0].text =	str(discipline1_3)
    v34	= doc.tables[2].cell(1,4).paragraphs[0].text =	str(discipline1_4)
    v35	= doc.tables[2].cell(1,5).paragraphs[0].text =	str(discipline1_5)
    v36	= doc.tables[2].cell(1,6).paragraphs[0].text =	str(discipline1_6)

    v37	= doc.tables[2].cell(2,1).paragraphs[0].text =	str(discipline2_1)
    v38	= doc.tables[2].cell(2,2).paragraphs[0].text =	str(discipline2_2)
    v39	= doc.tables[2].cell(2,3).paragraphs[0].text =	str(discipline2_3)
    v40	= doc.tables[2].cell(2,4).paragraphs[0].text =	str(discipline2_4)
    v41	= doc.tables[2].cell(2,5).paragraphs[0].text =	str(discipline2_5)
    v42	= doc.tables[2].cell(2,6).paragraphs[0].text =	str(discipline2_6)

    newdata.append(doc.paragraphs[4])
    newdata.append(doc.paragraphs[5])
    newdata.append(doc.tables[1].cell(1,1).paragraphs[0])
    newdata.append(doc.tables[1].cell(1,2).paragraphs[0])
    newdata.append(doc.tables[1].cell(1,3).paragraphs[0])
    newdata.append(doc.tables[1].cell(1,4).paragraphs[0])
    newdata.append(doc.tables[1].cell(1,5).paragraphs[0])
    newdata.append(doc.tables[1].cell(1,6).paragraphs[0])
    newdata.append(doc.tables[1].cell(1,7).paragraphs[0])

    newdata.append(doc.tables[1].cell(2,1).paragraphs[0])
    newdata.append(doc.tables[1].cell(2,2).paragraphs[0])
    newdata.append(doc.tables[1].cell(2,3).paragraphs[0])
    newdata.append(doc.tables[1].cell(2,4).paragraphs[0])
    newdata.append(doc.tables[1].cell(2,5).paragraphs[0])
    newdata.append(doc.tables[1].cell(2,6).paragraphs[0])
    newdata.append(doc.tables[1].cell(2,7).paragraphs[0])

    newdata.append(doc.tables[2].cell(1,1).paragraphs[0])
    newdata.append(doc.tables[2].cell(1,2).paragraphs[0])
    newdata.append(doc.tables[2].cell(1,3).paragraphs[0])
    newdata.append(doc.tables[2].cell(1,4).paragraphs[0])
    newdata.append(doc.tables[2].cell(1,5).paragraphs[0])
    newdata.append(doc.tables[2].cell(1,6).paragraphs[0])

    newdata.append(doc.tables[2].cell(2,1).paragraphs[0])
    newdata.append(doc.tables[2].cell(2,2).paragraphs[0])
    newdata.append(doc.tables[2].cell(2,3).paragraphs[0])
    newdata.append(doc.tables[2].cell(2,4).paragraphs[0])
    newdata.append(doc.tables[2].cell(2,5).paragraphs[0])
    newdata.append(doc.tables[2].cell(2,6).paragraphs[0])
    for data in newdata:
        #data.runs[0].underline = True
        #data.runs[0].underline = WD_UNDERLINE.THICK
        data.runs[0].font.name = 'TH SarabunIT๙'
        data.runs[0].font.size = Pt(16)

    doc.save(str(x+1)+'_'+name+'.docx')
    print("========End========")


